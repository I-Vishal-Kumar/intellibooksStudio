"""Document Processor with Parallel Processing for RAG Pipeline.

Supports multiple document types: PDF, DOCX, TXT, MD, HTML, etc.
Uses multiprocessing for parallel chunking and embedding generation.
"""

import os
import io
import hashlib
import logging
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from uuid import uuid4
import time

from pydantic import BaseModel

logger = logging.getLogger(__name__)

# Document loaders - lazy imports
_pypdf = None
_docx = None


def get_pypdf():
    global _pypdf
    if _pypdf is None:
        try:
            import pypdf
            _pypdf = pypdf
        except ImportError:
            logger.warning("pypdf not installed. PDF support unavailable.")
    return _pypdf


def get_docx():
    global _docx
    if _docx is None:
        try:
            import docx
            _docx = docx
        except ImportError:
            logger.warning("python-docx not installed. DOCX support unavailable.")
    return _docx


@dataclass
class ProcessingResult:
    """Result of document processing."""
    success: bool
    document_id: str
    filename: str
    chunks_created: int
    processing_time_ms: float
    error: Optional[str] = None
    metadata: Dict[str, Any] = None


class DocumentChunk(BaseModel):
    """A chunk of a document."""
    id: str
    content: str
    document_id: str
    chunk_index: int
    metadata: Dict[str, Any] = {}


class DocumentProcessor:
    """Process documents with parallel chunking for RAG."""

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.txt': 'text',
        '.md': 'markdown',
        '.html': 'html',
        '.htm': 'html',
        '.json': 'json',
        '.csv': 'csv',
    }

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        max_workers: int = 4,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_workers = max_workers

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        pypdf = get_pypdf()
        if pypdf is None:
            raise ImportError("pypdf not installed")

        text_parts = []
        reader = pypdf.PdfReader(io.BytesIO(content))

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX bytes."""
        docx_module = get_docx()
        if docx_module is None:
            raise ImportError("python-docx not installed")

        doc = docx_module.Document(io.BytesIO(content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    def _extract_text_from_html(self, content: bytes) -> str:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator='\n', strip=True)
        except ImportError:
            # Fallback: simple tag removal
            text = content.decode('utf-8', errors='ignore')
            import re
            text = re.sub(r'<[^>]+>', ' ', text)
            return ' '.join(text.split())

    def extract_text(self, content: bytes, file_extension: str) -> str:
        """Extract text from document based on type."""
        ext = file_extension.lower()

        if ext == '.pdf':
            return self._extract_text_from_pdf(content)
        elif ext in ['.docx', '.doc']:
            return self._extract_text_from_docx(content)
        elif ext in ['.html', '.htm']:
            return self._extract_text_from_html(content)
        elif ext in ['.txt', '.md', '.json', '.csv']:
            return content.decode('utf-8', errors='ignore')
        else:
            # Try to decode as text
            return content.decode('utf-8', errors='ignore')

    def _split_text_into_chunks(
        self,
        text: str,
        document_id: str,
        base_metadata: Dict[str, Any],
    ) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        chunks = []

        # Clean the text
        text = text.strip()
        if not text:
            return chunks

        # Split by sentences/paragraphs first for better chunk boundaries
        paragraphs = text.split('\n\n')

        current_chunk = ""
        chunk_index = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # If adding this paragraph exceeds chunk size, save current chunk
            if len(current_chunk) + len(para) + 2 > self.chunk_size and current_chunk:
                chunk_id = f"{document_id}_chunk_{chunk_index}"
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    content=current_chunk.strip(),
                    document_id=document_id,
                    chunk_index=chunk_index,
                    metadata={
                        **base_metadata,
                        "chunk_index": chunk_index,
                        "char_count": len(current_chunk),
                    }
                ))
                chunk_index += 1

                # Keep overlap from end of previous chunk
                if self.chunk_overlap > 0 and len(current_chunk) > self.chunk_overlap:
                    current_chunk = current_chunk[-self.chunk_overlap:]
                else:
                    current_chunk = ""

            current_chunk += ("\n\n" if current_chunk else "") + para

        # Add final chunk
        if current_chunk.strip():
            chunk_id = f"{document_id}_chunk_{chunk_index}"
            chunks.append(DocumentChunk(
                id=chunk_id,
                content=current_chunk.strip(),
                document_id=document_id,
                chunk_index=chunk_index,
                metadata={
                    **base_metadata,
                    "chunk_index": chunk_index,
                    "char_count": len(current_chunk),
                }
            ))

        return chunks

    def process_document(
        self,
        content: bytes,
        filename: str,
        document_id: Optional[str] = None,
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> Tuple[List[DocumentChunk], ProcessingResult]:
        """Process a single document into chunks.

        Returns tuple of (chunks, result).
        """
        start_time = time.time()

        # Generate document ID if not provided
        if document_id is None:
            content_hash = hashlib.md5(content).hexdigest()[:8]
            document_id = f"doc_{content_hash}_{uuid4().hex[:8]}"

        # Get file extension
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return [], ProcessingResult(
                success=False,
                document_id=document_id,
                filename=filename,
                chunks_created=0,
                processing_time_ms=0,
                error=f"Unsupported file type: {ext}",
            )

        try:
            # Extract text
            text = self.extract_text(content, ext)

            if not text.strip():
                return [], ProcessingResult(
                    success=False,
                    document_id=document_id,
                    filename=filename,
                    chunks_created=0,
                    processing_time_ms=(time.time() - start_time) * 1000,
                    error="No text content extracted",
                )

            # Base metadata
            base_metadata = {
                "document_id": document_id,
                "filename": filename,
                "file_type": self.SUPPORTED_EXTENSIONS[ext],
                "file_extension": ext,
                "total_chars": len(text),
                **(extra_metadata or {}),
            }

            # Split into chunks
            chunks = self._split_text_into_chunks(text, document_id, base_metadata)

            processing_time = (time.time() - start_time) * 1000

            return chunks, ProcessingResult(
                success=True,
                document_id=document_id,
                filename=filename,
                chunks_created=len(chunks),
                processing_time_ms=processing_time,
                metadata=base_metadata,
            )

        except Exception as e:
            logger.exception(f"Error processing document {filename}")
            return [], ProcessingResult(
                success=False,
                document_id=document_id,
                filename=filename,
                chunks_created=0,
                processing_time_ms=(time.time() - start_time) * 1000,
                error=str(e),
            )

    async def process_documents_parallel(
        self,
        documents: List[Tuple[bytes, str, Optional[str]]],  # (content, filename, doc_id)
    ) -> Tuple[List[DocumentChunk], List[ProcessingResult]]:
        """Process multiple documents in parallel.

        Args:
            documents: List of tuples (content_bytes, filename, optional_doc_id)

        Returns:
            Tuple of (all_chunks, all_results)
        """
        all_chunks = []
        all_results = []

        # Use ThreadPoolExecutor for I/O-bound operations
        loop = asyncio.get_event_loop()

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = []
            for content, filename, doc_id in documents:
                future = loop.run_in_executor(
                    executor,
                    self.process_document,
                    content,
                    filename,
                    doc_id,
                    None,
                )
                futures.append(future)

            # Gather all results
            results = await asyncio.gather(*futures)

            for chunks, result in results:
                all_chunks.extend(chunks)
                all_results.append(result)

        return all_chunks, all_results


class ParallelRAGProcessor:
    """High-level RAG processor with parallel document processing and indexing."""

    def __init__(
        self,
        vector_store,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        max_workers: int = 4,
        batch_size: int = 100,
    ):
        self.vector_store = vector_store
        self.document_processor = DocumentProcessor(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            max_workers=max_workers,
        )
        self.batch_size = batch_size

    async def ingest_document(
        self,
        content: bytes,
        filename: str,
        document_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessingResult:
        """Ingest a single document into the RAG system."""
        start_time = time.time()

        # Process document
        chunks, result = self.document_processor.process_document(
            content, filename, document_id, metadata
        )

        if not result.success or not chunks:
            return result

        # Convert chunks to vector store documents
        # Use the Document class from wherever available
        try:
            from .vector_store.chroma_store import Document
        except ImportError:
            from .vector_store.simple_store import Document

        documents = [
            Document(
                id=chunk.id,
                content=chunk.content,
                metadata=chunk.metadata,
            )
            for chunk in chunks
        ]

        # Index in batches for large documents
        for i in range(0, len(documents), self.batch_size):
            batch = documents[i:i + self.batch_size]
            await self.vector_store.add_documents(batch)

        # Update processing time to include indexing
        result.processing_time_ms = (time.time() - start_time) * 1000

        return result

    async def ingest_documents_parallel(
        self,
        documents: List[Tuple[bytes, str, Optional[str], Optional[Dict]]],
    ) -> List[ProcessingResult]:
        """Ingest multiple documents in parallel.

        Args:
            documents: List of tuples (content, filename, doc_id, metadata)
        """
        start_time = time.time()

        # Extract just the parts needed for processing
        docs_to_process = [(content, filename, doc_id) for content, filename, doc_id, _ in documents]

        # Process all documents in parallel
        all_chunks, results = await self.document_processor.process_documents_parallel(docs_to_process)

        if not all_chunks:
            return results

        # Convert chunks to vector store documents
        try:
            from .vector_store.chroma_store import Document
        except ImportError:
            from .vector_store.simple_store import Document

        vector_docs = [
            Document(
                id=chunk.id,
                content=chunk.content,
                metadata=chunk.metadata,
            )
            for chunk in all_chunks
        ]

        # Index all chunks in batches
        for i in range(0, len(vector_docs), self.batch_size):
            batch = vector_docs[i:i + self.batch_size]
            await self.vector_store.add_documents(batch)

        total_time = (time.time() - start_time) * 1000
        logger.info(
            f"Ingested {len(documents)} documents ({len(all_chunks)} chunks) in {total_time:.0f}ms"
        )

        return results

    async def delete_document(self, document_id: str) -> bool:
        """Delete a document and all its chunks."""
        return await self.vector_store.delete_by_metadata("document_id", document_id)

    async def query(
        self,
        query: str,
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ):
        """Query the RAG system."""
        return await self.vector_store.search(query, top_k, filters)

    async def get_stats(self) -> Dict[str, Any]:
        """Get RAG system statistics."""
        count = await self.vector_store.count()
        collections = await self.vector_store.list_collections()

        return {
            "total_chunks": count,
            "collections": collections,
            "chunk_size": self.document_processor.chunk_size,
            "chunk_overlap": self.document_processor.chunk_overlap,
        }
